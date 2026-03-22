from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"D:\workspace\EGS")
MD_PATH = ROOT / "caseC_report_draft.md"
DOCX_PATH = ROOT / "CaseC_Community_Microgrid_Report.docx"
OUTPUT_DIR = ROOT / "outputs_caseC"


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Case C: Community Microgrid\nCoursework Report")
    run.bold = True
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated from the verified modelling workflow in the EGS workspace.")
    run.italic = True
    run.font.size = Pt(10)


def add_markdown_content(document: Document, markdown_text: str) -> None:
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        p = document.add_paragraph()
        p.add_run(stripped)


def add_plot_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Selected Figures", level=1)

    figures = [
        ("Raw data sanity check", OUTPUT_DIR / "raw_data_sanity_check.png"),
        ("Base case: total load and PV", OUTPUT_DIR / "base_case_total_load_vs_pv.png"),
        ("Base case: battery SOC", OUTPUT_DIR / "base_case_soc.png"),
        ("Base case: grid import and export", OUTPUT_DIR / "base_case_grid_exchange.png"),
        ("Extension case: grid import and export", OUTPUT_DIR / "extension_case_grid_exchange.png"),
    ]

    for caption, image_path in figures:
        if not image_path.exists():
            continue
        document.add_heading(caption, level=2)
        document.add_picture(str(image_path), width=Inches(6.5))
        p = document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_document_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def main() -> None:
    markdown_text = MD_PATH.read_text(encoding="utf-8")

    document = Document()
    set_document_style(document)
    add_title(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_markdown_content(document, markdown_text)
    add_plot_section(document)
    document.save(DOCX_PATH)
    print(f"Saved Word report to: {DOCX_PATH}")


if __name__ == "__main__":
    main()
